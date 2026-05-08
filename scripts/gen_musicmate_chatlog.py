#!/usr/bin/env python3
"""Generate MusicMate-specific chat log Word document from all sessions."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime

# 關鍵字過濾
KEYWORDS = ['MusicMate', '練琴寶', 'musicmate', '練琴', 'OMR', '樂譜',
            '翻譜', '節拍器', '調音器', 'OpenCV', '翻頁']

doc = Document()

title = doc.add_heading('練琴寶 / MusicMate 開發聊天記錄', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.add_run('練琴寶專案討論記錄').bold = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph('※ 本檔案包含所有與練琴寶/MusicMate 專案相關的討論（自動從所有 session 過濾）')
doc.add_paragraph('※ 涵蓋 Telegram 和 Web UI 的完整對話')
doc.add_paragraph()

sessions_dir = '/home/aping/.openclaw/agents/main/sessions'
session_files = sorted([f for f in os.listdir(sessions_dir) if f.endswith('.jsonl')], reverse=True)

all_messages = []

def matches_project(text):
    text_lower = text.lower()
    for kw in KEYWORDS:
        if kw.lower() in text_lower:
            return True
    return False

for session_file in session_files:
    filepath = os.path.join(sessions_dir, session_file)
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                try:
                    record = json.loads(line.strip())
                    if record.get('type') != 'message':
                        continue

                    msg = record.get('message', {})
                    if msg.get('role') not in ('user', 'assistant'):
                        continue

                    text_content = []
                    if isinstance(msg.get('content'), list):
                        for c in msg['content']:
                            if c.get('type') == 'text':
                                text_content.append(c.get('text', ''))
                            elif c.get('type') == 'image':
                                text_content.append('[圖片]')
                    elif isinstance(msg.get('content'), str):
                        text_content.append(msg['content'])

                    text = '\n'.join(text_content).strip()
                    if not text:
                        continue

                    # 過濾：只保留與專案相關的訊息
                    if not matches_project(text):
                        continue

                    sender = msg.get('senderLabel', record.get('senderLabel', 'Unknown'))
                    timestamp = record.get('timestamp') or msg.get('timestamp', 0)
                    if timestamp:
                        try:
                            dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                            date_str = dt.strftime('%Y-%m-%d %H:%M')
                        except:
                            date_str = 'Unknown'
                    else:
                        date_str = 'Unknown'

                    role = 'William' if 'William' in str(sender) else '夥計'
                    all_messages.append((date_str, role, text))
                except Exception as e:
                    pass
    except Exception as e:
        pass

# Sort by timestamp
all_messages.sort(key=lambda x: x[0] if x[0] != 'Unknown' else '0000-00-00')

current_date = ""
for date_str, role, text in all_messages:
    date_only = date_str.split(' ')[0] if ' ' in date_str else date_str
    if date_only != current_date and date_only != 'Unknown':
        p = doc.add_paragraph()
        p.add_run(f"\n═══ 【{date_only}】 ═══").bold = True
        current_date = date_only

    p = doc.add_paragraph()
    p.add_run(f"[{date_str}] {role}：").bold = True
    p.add_run(text[:500] + ('...' if len(text) > 500 else ''))

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('※ 本記錄由 OpenClaw AI 自動生成').italic = True

output_path = '/home/aping/.openclaw/workspace/projects/musicmate/docs/Chat_Log_MusicMate.docx'
doc.save(output_path)
print(f'MusicMate chat log saved: {output_path}')
print(f'Total messages: {len(all_messages)}')
