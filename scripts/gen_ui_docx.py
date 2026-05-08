#!/usr/bin/env python3
"""Generate Word document for MusicMate UI Specification."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading('練琴寶 / MusicMate', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Windows 版 UI 規格書', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('版本：v0.1 Draft')
doc.add_paragraph('日期：2026-04-19')
doc.add_paragraph('負責：產品經理（AI 代理）')
doc.add_paragraph()

# Section 1
doc.add_heading('1. 整體架構', level=2)
doc.add_paragraph('採用單一主視窗設計，底部以 Tab 導航切換三大核心功能。')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('視窗屬性：').bold = True
doc.add_paragraph('• 最小尺寸：800 × 600 px', style='List Bullet')
doc.add_paragraph('• 預設尺寸：1200 × 800 px', style='List Bullet')
doc.add_paragraph('• 可縮放：是（響應式佈局）', style='List Bullet')
doc.add_paragraph('• 標題列：標準 Windows 標題列', style='List Bullet')
doc.add_paragraph('• 主題：跟隨 Windows 系統深色/淺色設定', style='List Bullet')

# Section 2
doc.add_heading('2. 主要功能區塊佈局', level=2)

# Table for layout
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.rows[0].cells[0]
cell.text = '''
┌─────────────────────────────────────────────────┐
│  標題列（系統標準）                              │
├─────────────────────────────────────────────────┤
│  頂部列：App 名稱 + 設定齒輪圖示                │
├─────────────────────────────────────────────────┤
│              主內容區（切換區）                  │
│   ┌─────────────────────────────────────────┐  │
│   │  自動翻譜  │  節拍器  │  調音器          │  │
│   └─────────────────────────────────────────┘  │
├─────────────────────────────────────────────────┤
│  底部 Tab 導航列（3 個 Tab 圖示 + 標籤）        │
└─────────────────────────────────────────────────┘
'''
doc.add_paragraph()

# Section 3
doc.add_heading('3. 自動翻譜頁（Sheet Music Auto-Scroll）', level=2)
doc.add_paragraph('用途：顯示樂譜 PDF/圖片，自動向上滾動，充當虛擬樂譜架。')

doc.add_heading('子區塊', level=3)
items = [
    ('樂譜顯示區', '顯示目前載入的 PDF 或圖片樂譜，支援滑鼠滾輪手動滾動，雙擊全螢幕模式'),
    ('控制工具列', '開啟樂譜、播放/暫停按鈕、速度滑桿（0.5x ~ 3.0x）、當前頁/總頁數'),
    ('頁面導航', '上一頁/下一頁 按鈕'),
    ('快捷鍵提示', 'Space=播放/暫停、←→=上一頁/下一頁、F=全螢幕'),
]
for name, desc in items:
    p = doc.add_paragraph()
    p.add_run(f'• {name}：').bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_heading('4. 節拍器頁（Metronome）', level=2)
doc.add_paragraph('用途：精準節拍器，協助練琴時維持穩定速度。')

items = [
    ('節拍顯示區', 'BPM 大字體、拍號、當前拍點閃爍動畫'),
    ('BPM 控制', '40~240 BPM、-10/-1/+1/+10 微調、快捷60/80/100/120'),
    ('拍號選擇', '2/4、3/4、4/4、5/4、6/8、7/8、12/8'),
    ('節拍音效', 'Click/Woodblock/Digital、音量控制'),
    ('計時器', '5/10/15/30/45/60 分鐘，倒數結束提醒'),
]
for name, desc in items:
    p = doc.add_paragraph()
    p.add_run(f'• {name}：').bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_heading('5. 調音器頁（Tuner）', level=2)
doc.add_paragraph('用途：收音麥克風偵測音高，協助樂器對音。')

items = [
    ('指針顯示', '半圓弧形指針、綠色=準確（±5 cents）、紅色=需調整'),
    ('音高顯示', '音名、Cents 偏差值、Hz 頻率'),
    ('麥克風控制', '啟動/停用，收音時圖示發綠光'),
    ('參考音設定', 'A4 頻率（預設 440 Hz，可調 430~450）、快捷432/440/442 Hz'),
    ('樂器預設', '吉他/烏克麗麗/小提琴/大提琴/長笛等'),
]
for name, desc in items:
    p = doc.add_paragraph()
    p.add_run(f'• {name}：').bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_heading('6. 設計語言', level=2)

table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Table Grid'
headers = [('項目', '說明'), ('主色調', '深藍 #0078D4（Fluent 藍）'), ('字體', 'Segoe UI Variable'), ('間距', '8px 基礎單位'), ('圓角', '4px（按鈕）、8px（卡片）、12px（面板）'), ('響應式', '主內容區採用彈性佈局')]
for i, (k, v) in enumerate(headers):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_heading('7. 預留未來功能', level=2)
doc.add_paragraph('• 音軌錄音', style='List Bullet')
doc.add_paragraph('• 練習日誌', style='List Bullet')
doc.add_paragraph('• 樂譜庫', style='List Bullet')
doc.add_paragraph('• 多視窗支援', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('本文件為 UI 規格初版，待 William 確認後可進入 Wireframe / Prototype 階段。').italic = True

output_path = '/home/aping/.openclaw/workspace/projects/musicmate/docs/MusicMate_UI_Spec_v0.1.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
