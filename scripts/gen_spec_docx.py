#!/usr/bin/env python3
"""Generate comprehensive Word document for MusicMate full specification v0.3."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 設定頁面邊距（單位：公分）
# 左、右、上、下 = 1.00 cm，裝訂邊 = 0.5 cm
section = doc.sections[0]
section.left_margin = Cm(1.00)
section.right_margin = Cm(1.00)
section.top_margin = Cm(1.00)
section.bottom_margin = Cm(1.00)
section.gutter = Cm(0.5)

# ==================== TITLE PAGE ====================
title = doc.add_heading('練琴寶 / MusicMate', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('完整產品規格書', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('版本：v1.1 Draft').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('日期：2026-04-29')
doc.add_paragraph('負責人：William（客人）/ 夥計（協調主持）')
doc.add_paragraph()
doc.add_paragraph('※ 本文件為 v1.1 完整規格書，已涵蓋所有討論過的功能與技術細節，含四個功能頁面完整 UI 示意圖。')
doc.add_page_break()

# ==================== SECTION 1: OVERVIEW ====================
doc.add_heading('1. 產品概述', level=1)
doc.add_heading('1.1 產品定位', level=2)
doc.add_paragraph('練琴寶（MusicMate）是一款面向 Windows 平台的練琴輔助工具，整合自動翻譜、節拍器、調音器與樂譜播放四大功能於單一應用程式中。')
doc.add_paragraph('產品目標：')
doc.add_paragraph('• 降低練琴時翻譜的干擾，讓練習更專注')
doc.add_paragraph('• 內建節拍器與調音器，減少額外硬體需求')
doc.add_paragraph('• 利用 MIDI 播放樂譜，提供聽覺參考輔助練習')
doc.add_paragraph('• 開源、免費、本地運行，保護用戶隱私')

doc.add_heading('1.2 目標用戶', level=2)
doc.add_paragraph('• 個人練習者（鋼琴、吉他、小提琴、長笛等）')
doc.add_paragraph('• 音樂教師（教學輔助、示範播放）')
doc.add_paragraph('• 音樂學生（考级練習、節奏訓練）')

doc.add_heading('1.3 產品特色', level=2)
doc.add_paragraph('• 四合一：翻譜 + 節拍器 + 調音器 + 播放，一應俱全')
doc.add_paragraph('• OMR 自動辨識：開啟 PDF 自動轉換為可播放結構')
doc.add_paragraph('• 本地運行：不依賴雲端，保護樂譜隱私')
doc.add_paragraph('• 跨樂器支援：涵蓋弦樂、管樂、鍵盤樂器')

doc.add_page_break()

# ==================== SECTION 2: FEATURES ====================
doc.add_heading('2. 產品功能列表', level=1)

doc.add_heading('2.1 自動翻譜功能', level=2)
features1 = [
    ('功能描述', '接收 PDF 或圖片格式的樂譜，自動顯示並根據樂曲速度自動滾動'),
    ('樂譜輸入', 'PDF 檔案開啟、圖片匯入（ JPG/PNG）、相機拍攝'),
    ('樂譜辨識', '使用 OMR（光學樂譜辨識）技術自動分析音符、節拍、速度'),
    ('顯示功能', '支援滑鼠滾輪手動滾動、雙擊全螢幕模式'),
    ('控制功能', '播放/暫停、上一頁/下一頁、速度調整（0.5x ~ 3.0x）'),
    ('快捷鍵', 'Space=播放/暫停、←→=換頁、F=全螢幕'),
    ('特殊功能', '無樂譜時顯示引導區塊、支援拖放檔案'),
]
table1 = doc.add_table(rows=len(features1), cols=2)
table1.style = 'Table Grid'
for i, (k, v) in enumerate(features1):
    table1.rows[i].cells[0].text = k
    table1.rows[i].cells[1].text = v

doc.add_heading('2.2 節拍器功能', level=2)
features2 = [
    ('功能描述', '提供精準節拍，協助練琴時維持穩定速度'),
    ('BPM 範圍', '40 ~ 240 BPM'),
    ('BPM 控制', '-10/-1/+1/+10 微調按鈕、快捷 60/80/100/120'),
    ('Tap Tempo', '使用者點擊或按鍵打拍子，自動計算並設定 BPM'),
    ('拍號選擇', '2/4、3/4、4/4、5/4、6/8、7/8、12/8'),
    ('節拍音效', '主機內建音效合成器為預設（Windows MIDI Synth / System Sounds），另備援預錄 WAV 樣本'),
    ('內建音色選項', 'Click（點擊）/ Woodblock（木魚）/ Digital（數位）/ Snare（小鼓）/ Bell（叮叮聲）/ Dog Bark（狗吠）'),
    ('音量控制', '獨立即量滑桿'),
    ('計時器', '5/10/15/30/45/60 分鐘倒數計時'),
    ('視覺效果', '節拍燈閃爍動畫'),
]
table2 = doc.add_table(rows=len(features2), cols=2)
table2.style = 'Table Grid'
for i, (k, v) in enumerate(features2):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v

doc.add_heading('2.3 調音器功能', level=2)
features3 = [
    ('功能描述', '收音麥克風偵測音高，協助樂器對音'),
    ('顯示方式', '半圓弧形指針儀表'),
    ('音高顯示', '音名（C/C#/D...）、Cents 偏差值、Hz 頻率'),
    ('精準範圍', '±5 cents 綠色區間 = 準確'),
    ('麥克風控制', '點擊啟動/停用，收音時顯示綠光指示'),
    ('參考音設定', 'A4 頻率（預設 440 Hz，可調 430~450 Hz）'),
    ('快捷參考音', '432 Hz / 440 Hz / 442 Hz'),
    ('樂器預設', '吉他/烏克麗麗/小提琴/大提琴/長笛等'),
    ('音高偵測演算法', '基於 FFT + 自相關（Autocorrelation）混合演算法，兼顧準確度與即時性'),
]
table3 = doc.add_table(rows=len(features3), cols=2)
table3.style = 'Table Grid'
for i, (k, v) in enumerate(features3):
    table3.rows[i].cells[0].text = k
    table3.rows[i].cells[1].text = v

doc.add_heading('2.4 樂譜播放功能（Audio Playback）【新】', level=2)
features4 = [
    ('功能描述', '使用電腦 MIDI 音源播放樂譜，邊聽邊練琴'),
    ('輸入來源', 'OMR 辨識後的音符結構'),
    ('MIDI 合成', '使用 General MIDI 內建模音源播放'),
    ('音色選擇', 'GM 標準音色（鋼琴、吉他、小提琴、長笛等）'),
    ('速度控制', '0.5x ~ 2.0x 播放速度調整'),
    ('連動功能', '邊播放邊自動翻譜（與自動翻譜頁連動）'),
    ('播放控制', '播放/暫停/停止、進度條顯示'),
    ('當前音符', '高亮顯示目前播放的音符位置'),
    ('循環播放（Loop）', '可設定小節範圍（如第 5-8 小節）反覆播放，便於分段練習'),
    ('移調（Transpose）', '可調整播放音高（-12 ~ +12 半音），便於不同調性練習'),
    ('分段練習', '可設定起始/結束小節，只播放特定段落'),
]
table4 = doc.add_table(rows=len(features4), cols=2)
table4.style = 'Table Grid'
for i, (k, v) in enumerate(features4):
    table4.rows[i].cells[0].text = k
    table4.rows[i].cells[1].text = v

doc.add_page_break()

# ==================== SECTION 3: TECH ARCHITECTURE ====================
doc.add_heading('3. 技術架構', level=1)

doc.add_heading('3.1 系統架構', level=2)
doc.add_paragraph('練琴寶採用三層架構設計：')
doc.add_paragraph('• 表現層（Presentation Layer）：Python GUI（PyQt/PySide）')
doc.add_paragraph('• 業務邏輯層（Business Logic）：核心功能模組')
doc.add_paragraph('• 資料處理層（Data Layer）：OMR 引擎、檔案處理')

doc.add_heading('3.2 OMR 光學樂譜辨識流程', level=2)
doc.add_paragraph('1. 樂譜圖片輸入（PDF 轉圖片 or 直接圖片）')
doc.add_paragraph('2. OpenCV 前處理：灰階轉換 → 降噪 → 傾斜校正 → 二值化')
doc.add_paragraph('3. 呼叫 Audiveris CLI 進行樂譜辨識')
doc.add_paragraph('4. 輸出 MusicXML 格式音符結構')
doc.add_paragraph('5. music21 解析 MusicXML，轉換為內部 JSON 結構')

doc.add_heading('3.3 MIDI 播放流程', level=2)
doc.add_paragraph('1. OMR 輸出（JSON 音符結構）→ mido 處理')
doc.add_paragraph('2. 轉換為 MIDI 訊息（Note On/Off、Program Change）')
doc.add_paragraph('3. pygame.simpleaudio 播放 MIDI 音頻')
doc.add_paragraph('4. 選擇 GM 音色（Program Change訊息）')
doc.add_paragraph('5. 速度調整（Playback Rate）')

doc.add_heading('3.4 主要技術堆疊', level=2)
tech_stack = [
    ('項目', '技術選擇'),
    ('程式語言', 'Python 3.12+'),
    ('GUI 框架', 'PyQt6（Windows Fluent Design）'),
    ('影像處理', 'OpenCV (cv2)'),
    ('OMR 辨識', 'Audiveris（開源 Java OMR 引擎）'),
    ('音頻處理', 'numpy、scipy（音訊分析）'),
    ('節拍器音效', '主機內建 MIDI Synthesizer / System Sounds（預設）+ 預錄 WAV 樣本（Click/Woodblock/Digital/Snare/Bell/DogBark）備援，pygame 播放'),
    ('MIDI 處理', 'mido（MIDI 訊息解析與生成）'),
    ('音訊播放', 'pygame / simpleaudio'),
    ('調音器收音', 'pyaudio / sounddevice（跨平台音訊擷取）'),
    ('PDF 處理', 'PyMuPDF（fitz）或 pdf2image'),
    ('MusicXML 解析', 'music21'),
    ('Word 輸出', 'python-docx（僅限文件生成工具）'),
]
table5 = doc.add_table(rows=len(tech_stack), cols=2)
table5.style = 'Table Grid'
for i, (k, v) in enumerate(tech_stack):
    table5.rows[i].cells[0].text = k
    table5.rows[i].cells[1].text = v

doc.add_page_break()

# ==================== SECTION 4: SYSTEM DATA FLOW ====================
doc.add_heading('4. 系統資料流程', level=1)

doc.add_heading('4.1 樂譜開啟流程', level=2)
doc.add_paragraph('1. 用戶點擊「開啟樂譜」→ 選擇 PDF/圖片檔案')
doc.add_paragraph('2. 檔案驗證：檢查檔案格式、大小、可讀取性')
doc.add_paragraph('3. PDF 轉圖片（PyMuPDF）→ 每頁轉為 300 DPI PNG')
doc.add_paragraph('4. 圖片前處理（OpenCV）→ 傾斜校正、降噪、二值化')
doc.add_paragraph('5. 顯示於「自動翻譜頁」')

doc.add_heading('4.2 OMR 辨識流程', level=2)
doc.add_paragraph('1. 用戶點擊「辨識樂譜」按鈕')
doc.add_paragraph('2. 呼叫 Audiveris CLI（subprocess）→ 輸入圖片')
doc.add_paragraph('3. Audiveris 輸出 MusicXML')
doc.add_paragraph('4. music21 解析 MusicXML → 轉換為內部 JSON 結構')
doc.add_paragraph('5. JSON 結構供「樂譜播放頁」使用')

doc.add_heading('4.3 播放流程', level=2)
doc.add_paragraph('1. 用戶啟動播放 → 讀取 JSON 音符結構')
doc.add_paragraph('2. mido 將 JSON 轉換為 MIDI 訊息（Note On/Off、Program Change）')
doc.add_paragraph('3. pygame 播放 MIDI 音頻 → 同步高亮顯示當前音符')
doc.add_paragraph('4. 自動翻譜頁根據播放進度自動滾動樂譜')

doc.add_heading('4.4 調音器流程', level=2)
doc.add_paragraph('1. 用戶點擊麥克風按鈕 → 啟動 pyaudio 收音串流')
doc.add_paragraph('2. FFT + 自相關演算法即時分析音高')
doc.add_paragraph('3. UI 更新指針、音名、Cents 偏差值、Hz 頻率')
doc.add_paragraph('4. 停用麥克風 → 關閉收音串流，釋放資源')

doc.add_heading('4.5 設定儲存流程', level=2)
doc.add_paragraph('1. 用戶修改設定（BPM、音量、主題等）')
doc.add_paragraph('2. 設定寫入 %APPDATA%/MusicMate/config/settings.json')
doc.add_paragraph('3. 應用程式啟動時讀取設定檔，套用上次設定')

doc.add_heading('4.6 節拍器流程', level=2)
doc.add_paragraph('1. 用戶選擇節拍音效與 BPM/拍號')
doc.add_paragraph('2. 系統依據 BPM 計算每拍之間的毫秒間隔（60000 / BPM ms）')
doc.add_paragraph('3. 根據拍號，系統計算每小節的拍數（例如 4/4 = 4 拍，3/4 = 3 拍）')
doc.add_paragraph('4. 播放時：')
doc.add_paragraph('   a. 強拍（第一拍）：使用較強音效（Main Sound）')
doc.add_paragraph('   b. 弱拍（其餘拍）：使用較輕音效（Secondary Sound）')
doc.add_paragraph('   c. 每到達一拍時，同時更新 BPM 顯示、拍號位置與閃爍動畫')
doc.add_paragraph('5. 音效來源（優先順序）：')
doc.add_paragraph('   a. 主機內建音效合成器（Windows MIDI Synthesizer / System Sounds API）— 預設使用此方式')
doc.add_paragraph('   b. 備援：預錄 WAV 樣本（Click/Woodblock/Digital）由 pygame 播放')
doc.add_paragraph('6. Tap Tempo：用戶連續點擊，系統統計點擊間隔計算 BPM（取最近 4 次平均）')
doc.add_paragraph('7. 計時器倒數完畢，自動停止節拍器並播放提示音')

doc.add_page_break()

# ==================== SECTION 5: UI ====================
doc.add_heading('5. Windows 版 UI 規格', level=1)

doc.add_heading('5.1 整體架構', level=2)
doc.add_paragraph('• 視窗模型：單一主視窗 + 底部 Tab 導航（4個 Tab）')
doc.add_paragraph('• 最小尺寸：800 × 600 px')
doc.add_paragraph('• 預設尺寸：1200 × 800 px')
doc.add_paragraph('• 可縮放：是（響應式佈局）')
doc.add_paragraph('• 標題列：標準 Windows 標題列')
doc.add_paragraph('• 主題：跟隨 Windows 系統深色/淺色設定')

doc.add_heading('5.2 佈局結構', level=2)
doc.add_paragraph('單一視窗分為以下四個區塊（由上而下）：')
doc.add_paragraph()

# 插入 UI 佈局示意圖
diagram_path = '/home/aping/.openclaw/workspace/projects/musicmate/docs/ui_layout_diagram.png'
doc.add_picture(diagram_path, width=Inches(6.0))
doc.add_paragraph()
doc.add_paragraph('▲ 圖：練琴寶 MusicMate Windows 版 UI 佈局結構')

doc.add_heading('5.3 四大功能頁面', level=2)

# Page 1: Auto-scroll
auto_scroll_img = '/home/aping/.openclaw/media/tool-image-generation/gui_auto_scroll---e8f4ab6c-b2db-4626-9fec-81a8b1dc38a0.png'
doc.add_heading('5.3.1 自動翻譜頁', level=3)
doc.add_paragraph('用途：顯示樂譜 PDF/圖片，自動向上滾動，充當虛擬樂譜架。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 樂譜顯示區：PDF/圖片顯示、支援滑鼠滾輪、雙擊全螢幕')
doc.add_paragraph('• 控制工具列：開啟樂譜、播放/暫停、速度滑桿（0.5x~3.0x）')
doc.add_paragraph('• 頁面導航：上一頁/下一頁按鈕')
doc.add_paragraph('• 快捷鍵：Space=播放/暫停、←→=換頁、F=全螢幕')
doc.add_paragraph()
doc.add_picture(auto_scroll_img, width=Inches(6.0))
doc.add_paragraph('▲ 圖：自動翻譜頁 UI 示意')
doc.add_paragraph()

# Page 2: Metronome
metronome_img = '/home/aping/.openclaw/media/tool-image-generation/gui_metronome---c70feaf4-9d53-4ff2-8fb4-89718fd72cef.png'
doc.add_heading('5.3.2 節拍器頁', level=3)
doc.add_paragraph('用途：精準節拍器，協助練琴時維持穩定速度。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 節拍顯示：BPM 大字體、拍號、閃爍動畫')
doc.add_paragraph('• BPM 控制：40~240、快捷 60/80/100/120')
doc.add_paragraph('• 拍號選擇：2/4、3/4、4/4、5/4、6/8、7/8、12/8')
doc.add_paragraph('• 音效：Click（點擊）/ Woodblock（木魚）/ Digital（數位）/ Snare（小鼓）/ Bell（叮叮聲）/ Dog Bark（狗吠），使用主機內建音效合成器播放，WAV 樣本備援')
doc.add_paragraph('• 計時器：5~60 分鐘倒數')
doc.add_paragraph()
doc.add_picture(metronome_img, width=Inches(6.0))
doc.add_paragraph('▲ 圖：節拍器頁 UI 示意')
doc.add_paragraph()

# Page 3: Tuner
tuner_img = '/home/aping/.openclaw/media/tool-image-generation/gui_tuner---0b9e621a-ed89-4197-856d-b07aaa7cd08e.png'
doc.add_heading('5.3.3 調音器頁', level=3)
doc.add_paragraph('用途：收音麥克風偵測音高，協助樂器對音。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 指針顯示：半圓弧形，±5 cents 綠色=準')
doc.add_paragraph('• 音高顯示：音名、Cents、Hz')
doc.add_paragraph('• 麥克風：啟動/停用控制')
doc.add_paragraph('• 參考音：A4 預設 440 Hz（古典音樂常用 443 Hz，歷史演奏常用 432 Hz），可調範圍 430~450 Hz，432/440/442 Hz 快捷鍵')
doc.add_paragraph('• 樂器預設：吉他/小提琴/長笛等')
doc.add_paragraph()
doc.add_picture(tuner_img, width=Inches(6.0))
doc.add_paragraph('▲ 圖：調音器頁 UI 示意')
doc.add_paragraph()

# Page 4: Audio Playback
playback_img = '/home/aping/.openclaw/media/tool-image-generation/gui_playback---d42f1a3f-99f0-47aa-92c7-3f442132c133.png'
doc.add_heading('5.3.4 樂譜播放頁（Audio Playback）【新】', level=3)
doc.add_paragraph('用途：使用電腦 MIDI 音源播放樂譜，邊聽邊練琴。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 播放控制：播放/暫停、停止按鈕')
doc.add_paragraph('• 速度控制：滑桿 0.5x ~ 2.0x')
doc.add_paragraph('• 音色選擇：下拉選單（鋼琴/吉他/小提琴等 GM 音色）')
doc.add_paragraph('• 音量控制：獨立即量滑桿')
doc.add_paragraph('• 進度條：顯示目前播放位置')
doc.add_paragraph('• 當前音符：高亮顯示目前播放的音符')
doc.add_paragraph('• 音軌選擇：多音軌樂譜（鋼琴左右手、弦樂重奏等）可選擇播放特定音軌或全部混音')
doc.add_paragraph()
doc.add_picture(playback_img, width=Inches(6.0))
doc.add_paragraph('▲ 圖：樂譜播放頁 UI 示意')
doc.add_paragraph()

doc.add_heading('5.3.4.1 Soundfont 音源方案', level=3)
doc.add_paragraph('• 第一版使用 Windows 作業系統內建 General MIDI Soundfont（gs soft synth）')
doc.add_paragraph('• 若內建音源品質不足，預設使用 PyGame 內建模音源（ fluidsynth ）')
doc.add_paragraph('• 音色採用 GM 標準映射，Program Change 訊息控制音色切換')
doc.add_paragraph('• 安裝包不捆绑額外 Soundfont 檔案，未來可擴充使用者自訂 Soundfont 功能')

doc.add_heading('5.3.4.2 自動翻譜同步機制', level=3)
doc.add_paragraph('播放進度與自動翻譜的同步邏輯：')
doc.add_paragraph('• 每個音符帶有 page（頁數）與 measure（小節）屬性，播放時根據目前時間戳記計算所在小節')
doc.add_paragraph('• 跨頁長音符（如鋼琴踏板維持超過一頁）：系統記錄音符起始頁，跨越頁面邊界時不額外翻頁')
doc.add_paragraph('• 若某頁無音符（小節空白），該頁不主動翻頁，等待下一頁有音符時再翻頁')
doc.add_paragraph('• 自動翻譜預設落後播放進度一個小節，避免視覺領先聽覺造成不協調')
doc.add_paragraph('• 使用者可在設定中調整「翻頁提前量」（0~2 小節）')

doc.add_heading('5.3.5 連動功能說明', level=3)
doc.add_paragraph('自動翻譜頁 + 樂譜播放頁可同時運行：')
doc.add_paragraph('• 啟動播放時，樂譜自動跟著滾動')
doc.add_paragraph('• 節拍器可與播放同步')
doc.add_paragraph('• 可隨時切換手動/自動模式')

doc.add_heading('5.4 檔案選單', level=2)
doc.add_paragraph('• 開啟樂譜（Ctrl+O）：支援 PDF、JPG、PNG 格式')
doc.add_paragraph('• 最近開啟的檔案（Recent Files）：記憶最近 10 個樂譜，快速存取')
doc.add_paragraph('• 另存為：將當前處理結果匯出為 MusicXML 或 PDF')
doc.add_paragraph('• 離開（Alt+F4）：關閉應用程式，自動儲存設定')

doc.add_heading('5.5 拖放支援', level=2)
doc.add_paragraph('• 支援將 PDF/圖片檔案直接拖曳至主視窗')
doc.add_paragraph('• 拖曳至「自動翻譜頁」：立即載入樂譜')
doc.add_paragraph('• 拖曳至非翻譜頁：自動切換至翻譜頁並載入')
doc.add_paragraph('• 拖曳多個檔案：依序加入最近檔案清單，載入第一個')

doc.add_heading('5.6 設定對話框', level=2)
doc.add_paragraph('• 存取方式：選單 → 設定（Ctrl+,）或 Tab 列齒輪圖示')
doc.add_paragraph('• 分頁內容：')
doc.add_paragraph('  - 一般：語言（繁體中文/English）、主題（淺色/深色/跟隨系統）、啟動時檢查更新')
doc.add_paragraph('  - OMR：解析度設定（150/200/300 DPI）、辨識超時秒數（預設 60s）')
doc.add_paragraph('  - 播放：預設音色、預設速度、循環播放預設設定')
doc.add_paragraph('  - 節拍器：預設 BPM、預設拍號、預設音效')
doc.add_paragraph('  - 調音器：預設 A4 頻率、預設樂器')
doc.add_paragraph('  - 快捷鍵：自訂鍵盤快捷鍵（可檢視目前對照表）')

doc.add_heading('5.7 鍵盤快捷鍵對照表', level=2)
shortcuts = [
    ('快捷鍵', '功能'),
    ('Ctrl+O', '開啟樂譜'),
    ('Ctrl+,', '開啟設定對話框'),
    ('Ctrl+W', '關閉當前樂譜'),
    ('Space', '播放/暫停'),
    ('←/→', '上一頁/下一頁'),
    ('F', '全螢幕切換'),
    ('Tab', '在 Tab 之間切換（翻譜/節拍器/調音器/播放）'),
    ('1/2/3/4', '直接切換至對應 Tab'),
    ('M', '節拍器啟動/停止'),
    ('T', '調音器麥克風啟動/停用'),
    ('L', '循環播放開關'),
    ('+ / -（主鍵盤）', 'BPM +1/-1'),
    ('Page Up/Page Down', '翻譜上一頁/下一頁'),
    ('Esc', '退出全螢幕'),
    ('Alt+F4', '離開應用程式'),
]
table_shortcuts = doc.add_table(rows=len(shortcuts), cols=2)
table_shortcuts.style = 'Table Grid'
for i, row_data in enumerate(shortcuts):
    for j, cell_data in enumerate(row_data):
        table_shortcuts.rows[i].cells[j].text = cell_data

doc.add_heading('5.8 關於對話框', level=2)
doc.add_paragraph('• 存取方式：選單 → 說明 → 關於')
doc.add_paragraph('• 顯示內容：')
doc.add_paragraph('  - 產品名稱：練琴寶 MusicMate')
doc.add_paragraph('  - 版本號（v1.0.0）')
doc.add_paragraph('  - 版權宣告：© 2026 William Chang')
doc.add_paragraph('  - 授權條款：MIT License')
doc.add_paragraph('  - 開源元件清單（PyQt6、Audiveris、OpenCV 等，含連結）')
doc.add_paragraph('  - 回報問題按鈕（連結至 GitHub Issues）')

doc.add_page_break()

# ==================== SECTION 6: OMR ====================
doc.add_heading('6. 光學樂譜辨識（OMR）方案', level=1)

doc.add_heading('6.1 為何使用 OpenCV？', level=2)
doc.add_paragraph('• Python 原生，跨平台方便')
doc.add_paragraph('• 影像前處理功能強大（降噪、傾斜校正、二值化）')
doc.add_paragraph('• 生態系完整，搭配其他函式庫容易')

doc.add_heading('6.2 OMR 處理流程（Prototype → Audiveris）', level=2)
doc.add_paragraph('第一版採用 Audiveris（開源 OMR）做 Prototype，流程如下：')
doc.add_paragraph('Step 1：樂譜輸入（PDF 轉圖片 or 直接圖片）')
doc.add_paragraph('Step 2：PDF 處理（PyMuPDF 轉圖片）')
doc.add_paragraph('Step 3：呼叫 Audiveris 命令列介面（CLI）進行樂譜辨識')
doc.add_paragraph('Step 4：輸出 MusicXML 格式音符結構')
doc.add_paragraph('Step 5：解析 MusicXML，轉換為內部 JSON 結構')

doc.add_heading('6.3 Audiveris 整合方案', level=2)
doc.add_paragraph('Audiveris 是目前最成熟的開源 OMR，採用以下整合方式：')
doc.add_paragraph('• 使用 N4 (N Audrey) 版本，支援命令列批次處理')
doc.add_paragraph('• Java 環境：需安裝 JRE 17+（Windows 11 環境）')
doc.add_paragraph('• Python 整合：subprocess 呼叫 Audiveris CLI，讀取 MusicXML 輸出')
doc.add_paragraph('• PDF 轉圖片：使用 PyMuPDF（fitz）處理 PDF 頁面轉圖片')
doc.add_paragraph('• 支援格式：PDF / JPG / PNG / TIFF 輸入')
doc.add_paragraph('• 輸出格式：MusicXML（標準音樂交換格式）')

doc.add_heading('6.4 技術堆疊（OMR 部分）', level=2)
omr_stack = [
    ('項目', '技術選擇'),
    ('PDF 處理', 'PyMuPDF（fitz）'),
    ('OMR 引擎', 'Audiveris（開源，Java）'),
    ('OMR 前處理', 'OpenCV 輔助（傾斜校正、降噪）'),
    ('音樂格式解析', 'music21（讀取 MusicXML）'),
    ('內部資料結構', 'JSON（音符、時值、節拍、小節）'),
]
table_omr = doc.add_table(rows=len(omr_stack), cols=2)
table_omr.style = 'Table Grid'
for i, (k, v) in enumerate(omr_stack):
    table_omr.rows[i].cells[0].text = k
    table_omr.rows[i].cells[1].text = v

doc.add_heading('6.5 Phase 5 規劃（未來自訓練模型）', level=2)
doc.add_paragraph('第一版使用 Audiveris 快速驗證可行性。\n未來若要提升準確度、支援手寫樂譜，可考虑：')
doc.add_paragraph('• 採用公開資料集訓練（MUSCIMA++、HASYTAC）')
doc.add_paragraph('• 模型架構：CNN + CTC（主流 OMR 做法）')
doc.add_paragraph('• 評估時機：Audiveris 準確度瓶頸出现後再決定')
doc.add_paragraph('• 優先序：★★★☆☆（目前擺後面）')

doc.add_page_break()

# ==================== SECTION 7: PDF OMR PROCESS ====================
doc.add_heading('7. PDF 樂譜 OMR 處理流程（第一版）', level=1)

doc.add_heading('7.1 整體流程', level=2)
doc.add_paragraph('PDF 格式樂譜的 OMR 處理分為以下 6 個步驟：')
process_steps = [
    ('步驟', '內容', '技術工具'),
    ('Step 1', '使用者開啟 PDF 樂譜', 'PyQt6 檔案對話框'),
    ('Step 2', 'PDF 轉圖片（每頁一張）', 'PyMuPDF（fitz）'),
    ('Step 3', '圖片前處理（傾斜/降噪/二值化）', 'OpenCV'),
    ('Step 4', 'Audiveris OMR 辨識', 'subprocess 呼叫 Audiveris CLI'),
    ('Step 5', '輸出 MusicXML 檔案', 'Audiveris'),
    ('Step 6', '解析 MusicXML 轉 JSON', 'music21'),
]
table_proc = doc.add_table(rows=len(process_steps), cols=3)
table_proc.style = 'Table Grid'
for i, row_data in enumerate(process_steps):
    for j, cell_data in enumerate(row_data):
        table_proc.rows[i].cells[j].text = cell_data

doc.add_heading('7.2 詳細步驟說明', level=2)
doc.add_paragraph('【Step 1】使用者開啟 PDF 樂譜')
doc.add_paragraph('• 使用者點擊「開啟樂譜」按鈕，選擇 PDF 檔案')
doc.add_paragraph('• 練琴寶檢查檔案是否存在，記錄檔案路徑')
doc.add_paragraph('• 顯示載入進度條（PDF 頁數可能達到數十頁）')

doc.add_paragraph('【Step 2】PDF 轉圖片（PyMuPDF）')
doc.add_paragraph('• 使用 PyMuPDF（fitz）將 PDF 每頁轉為圖片')
doc.add_paragraph('• 建議解析度：300 DPI（印刷品質）')
doc.add_paragraph('• 圖片格式：PNG（無失真，適合 OMR）')
doc.add_paragraph('• 支援多頁同時處理（批次轉換）')
doc.add_paragraph('• 臨時圖片存放在記憶體或暫存資料夾，處理完畢後自動刪除')

doc.add_paragraph('【Step 3】圖片前處理（OpenCV）')
doc.add_paragraph('• 傾斜校正（Deskew）：偵測水平線，自動旋轉')
doc.add_paragraph('• 灰階轉換（Grayscale）：降低運算量')
doc.add_paragraph('• 對比度增強（Contrast）：讓音符更明顯')
doc.add_paragraph('• 降噪（Denoise）：移除灰塵/掃描雜點')
doc.add_paragraph('• 二值化（Binarize）：轉為黑白，提高辨識度')

doc.add_paragraph('【Step 4】Audiveris OMR 辨識')
doc.add_paragraph('• 呼叫 Audiveris N4 命令列介面（CLI）')
doc.add_paragraph('• 批次處理多張圖片')
doc.add_paragraph('• Java 環境需 JRE 17+（Windows 11 環境）')
doc.add_paragraph('• 處理時間視 PDF 頁數而定（約數秒至數十秒）')

doc.add_paragraph('【Step 5】輸出 MusicXML 檔案')
doc.add_paragraph('• Audiveris 輸出標準 MusicXML 格式')
doc.add_paragraph('• 包含完整音符、時值、節拍、小節、調號資訊')
doc.add_paragraph('• 樂器軌道、動態、歌詞等額外資訊（可選）')

doc.add_paragraph('【Step 6】解析 MusicXML 轉 JSON（music21）')
doc.add_paragraph('• 使用 music21 函式庫讀取 MusicXML')
doc.add_paragraph('• 轉換為練琴寶內部 JSON 結構')
doc.add_paragraph('• 結構包含：頁數、tempo、拍號、每頁音符陣列')

doc.add_heading('7.3 輸出 JSON 結構範例', level=2)
doc.add_paragraph('{')
doc.add_paragraph('  "pages": [{"pageIndex": 0, "width": 800, "height": 1100}],')
doc.add_paragraph('  "tempo": 120,')
doc.add_paragraph('  "timeSignature": "4/4",')
doc.add_paragraph('  "keySignature": "C major",')
doc.add_paragraph('  "notes": [')
doc.add_paragraph('    {"page": 1, "x": 100, "y": 200, "pitch": "C4", "duration": "quarter", "measure": 1},')
doc.add_paragraph('    {"page": 1, "x": 150, "y": 200, "pitch": "D4", "duration": "quarter", "measure": 1}')
doc.add_paragraph('  ]')
doc.add_paragraph('}')

doc.add_page_break()

# ==================== SECTION 8: NON-FUNCTIONAL REQUIREMENTS ====================
doc.add_heading('8. 非功能需求', level=1)

doc.add_heading('8.1 效能需求', level=2)
doc.add_paragraph('• PDF 載入時間：不超過 3 秒（10 頁以內）')
doc.add_paragraph('• OMR 辨識速度：每頁不超過 5 秒（Audiveris CLI）')
doc.add_paragraph('• MIDI 播放延遲：小於 50ms')
doc.add_paragraph('• UI 回應時間：所有操作在 100ms 內完成')
doc.add_paragraph('• 記憶體使用：不超過 500MB（一般樂譜）')
doc.add_paragraph('• 啟動時間：不超過 5 秒')

doc.add_heading('8.2 相容性', level=2)
doc.add_paragraph('作業系統：')
doc.add_paragraph('• Windows 10 (64-bit) 及以上版本')
doc.add_paragraph('• Windows 11 完整支援')
doc.add_paragraph('Python 環境：')
doc.add_paragraph('• Python 3.12 及以上版本')
doc.add_paragraph('外部相依：')
doc.add_paragraph('• Java JRE 17+（Audiveris 運行環境，詳見 Section 9.3.2 完整套件清單）')
doc.add_paragraph('螢幕解析度：最小 800x600 px，建議 1280x720 px 及以上')

doc.add_heading('8.3 可靠度', level=2)
doc.add_paragraph('• 連續運行 8 小時無崩潰')
doc.add_paragraph('• 所有外部操作需有錯誤處理機制')
doc.add_paragraph('• 日誌記錄：所有關鍵操作需記錄至 log 檔')

doc.add_heading('8.4 可用性', level=2)
doc.add_paragraph('• 新手引導：首次使用時顯示功能說明')
doc.add_paragraph('• 快速鍵：所有常用功能需支援鍵盤快捷鍵')
doc.add_paragraph('• 多語言：預設繁體中文，未來支援英文')
doc.add_paragraph('• 無障礙（Accessibility）：')
doc.add_paragraph('  - 字型縮放：UI 字型可放大至 150%')
doc.add_paragraph('  - 高對比模式：提供高對比主題選項')
doc.add_paragraph('  - 鍵盤導航：所有功能可透過鍵盤操作，不需滑鼠')
doc.add_paragraph('  - 螢幕閱讀器：支援 Windows Narrator 基本語音提示')

doc.add_heading('8.4.1 測試策略', level=2)
doc.add_paragraph('• 單元測試：所有核心模組（節拍器、調音器、MIDI、OMR）需有單元測試，採用 pytest 框架')
doc.add_paragraph('• 測試覆蓋率目標：關鍵模組覆蓋率 > 80%')
doc.add_paragraph('• 整合測試：各模組之間的介面互動需有整合測試')
doc.add_paragraph('• UI 測試：PyQt6 測試採用 pytest-qt，測試主要使用者流程')
doc.add_paragraph('• 自動化測試：CI/CD 流程中自動執行所有測試（見 Section 9.3.8）')

doc.add_heading('8.5 安全性', level=2)
doc.add_paragraph('• 本地處理：所有樂譜處理均在本地完成，不上传至雲端')
doc.add_paragraph('• 隱私保護：不收集用戶個資或使用習慣')
doc.add_paragraph('• 輸入驗證：所有使用者輸入（檔案路徑、設定值）需驗證與過濾')

doc.add_heading('8.6 擴展性', level=2)
doc.add_paragraph('• 插件系統：預留插件介面，未來可擴充功能')
doc.add_paragraph('• 模組化設計：各功能模組獨立，方便替換或升級')
doc.add_paragraph('• API 文檔：各模組提供清晰的介面定義與使用範例')

doc.add_heading('8.7 資料持久化', level=2)
doc.add_paragraph('• 設定檔：採用 JSON 格式，存放於 %APPDATA%/MusicMate/config/settings.json')
doc.add_paragraph('• 最近檔案清單：記錄於 %APPDATA%/MusicMate/config/recent_files.json（最多 10 筆）')
doc.add_paragraph('• OMR 快取：辨識結果快取於 %APPDATA%/MusicMate/cache/，避免重複辨識相同檔案')
doc.add_paragraph('• 日誌：存放於 %APPDATA%/MusicMate/logs/，參見 Section 9.1.1')
doc.add_paragraph('• 首次執行：自動建立上述目錄與預設設定檔')

doc.add_heading('8.8 檔案關聯', level=2)
doc.add_paragraph('• 可選註冊為 .pdf 和 .musicxml 檔案的預設開啟程式之一')
doc.add_paragraph('• 安裝時提供選項，不強制設定為預設程式')
doc.add_paragraph('• 支援雙擊檔案直接開啟，自動載入對應功能頁面')

doc.add_page_break()

# ==================== SECTION 9: ERROR HANDLING & INSTALLATION ====================
doc.add_heading('9. 錯誤處理與安裝發布', level=1)

doc.add_heading('9.1 錯誤處理機制', level=2)
doc.add_paragraph('練琴寶採用分層錯誤處理策略：')
doc.add_paragraph('UI 層：所有使用者操作需有 try-except 保護，錯誤訊息需顯示原因與解決建議')
doc.add_paragraph('業務邏輯層：驗證所有輸入參數，使用日誌模組（logging）記錄詳細錯誤資訊')
doc.add_paragraph('資料處理層：檔案操作需檢查檔案是否存在、權限是否足夠，外部工具失敗需有重試機制（最多 3 次）')

doc.add_heading('9.1.1 日誌規範', level=3)
doc.add_paragraph('• 日誌層級：DEBUG（開發）、INFO（一般）、WARNING（可復原異常）、ERROR（操作失敗）、CRITICAL（系統崩潰）')
doc.add_paragraph('• 日誌格式：[YYYY-MM-DD HH:MM:SS] [LEVEL] [模組] 訊息')
doc.add_paragraph('• 存放位置：%APPDATA%/MusicMate/logs/musicmate.log')
doc.add_paragraph('• 日誌輪轉：單一檔案上限 10MB，保留最近 5 個備份，自動刪除最舊檔案')

doc.add_heading('9.1.2 自訂 Exception 類別', level=3)
doc.add_paragraph('定義專案專用例外類別，便於分層捕捉與處理：')
custom_exc = [
    ('例外類別', '觸發時機', '處理方式'),
    ('PDFLoadError', 'PDF 檔案損壞、格式不支援、無法讀取', 'UI 顯示錯誤訊息，記錄日誌'),
    ('OMRError', 'Audiveris 辨識失敗、輸出 MusicXML 損壞', '提示用戶重試或檢查圖片品質'),
    ('AudioDeviceError', '麥克風/音效設備無法存取、被其他程式佔用', '引導用戶檢查裝置設定'),
    ('MIDIError', 'MIDI 設備未設定、Pygame 初始化失敗', '引導至 Windows 音效設定'),
    ('JavaRuntimeError', '系統未安裝 Java JRE 或版本不符', '提示下載並安裝 Java JRE 17+'),
    ('SubprocessTimeoutError', 'Audiveris 執行超時（預設 60 秒）', '終止程序，提示用戶稍後重試'),
]
table_exc = doc.add_table(rows=len(custom_exc), cols=3)
table_exc.style = 'Table Grid'
for i, row_data in enumerate(custom_exc):
    for j, cell_data in enumerate(row_data):
        table_exc.rows[i].cells[j].text = cell_data

doc.add_heading('9.1.3 Audiveris 子程序崩潰處理', level=3)
doc.add_paragraph('• 執行 Audiveris CLI 時，若 Java 程序異常退出（exit code != 0），自動記錄輸出與錯誤串流')
doc.add_paragraph('• 重試機制：最多重試 2 次，間隔 3 秒')
doc.add_paragraph('• 若所有重試失敗，保留暫存圖片供除錯，並顯示「OMR 引擎異常」對話框，附帶錯誤碼與回報按鈕')
doc.add_paragraph('• 超時保護：單一頁面處理上限 60 秒，逾時強制終止子程序')

doc.add_heading('9.1.4 錯誤對話框規範', level=3)
doc.add_paragraph('• 樣式：使用 PyQt6 QMessageBox 標準對話框，包含圖示（警告/錯誤/資訊）')
doc.add_paragraph('• 內容結構：標題（簡短描述）+ 內文（原因說明 + 解決建議）+ 操作按鈕（重試 / 取消 / 回報問題）')
doc.add_paragraph('• 錯誤碼：每種錯誤情境分配唯一代碼（如 ERR-PDF-001），便於除錯追蹤')
doc.add_paragraph('• 「回報問題」按鈕：自動打包最近 50 行日誌，提供用戶可選擇是否附加當前樂譜檔案')

doc.add_heading('9.2 常見錯誤與解決方案', level=2)
error_handling = [
    ('錯誤情境', '錯誤碼', '可能原因', '解決方案'),
    ('PDF 無法開啟', 'ERR-PDF-001', '檔案損壞或格式不支援', '顯示「無法開啟此檔案，請確認檔案未損壞」'),
    ('OMR 辨識結果空白', 'ERR-OMR-001', '圖片解析度過低或頁面為純文字', '提示使用者確認上傳的是掃描圖片而非文字 PDF'),
    ('Audiveris 執行超時', 'ERR-OMR-002', '頁面過大或系統資源不足', '提示用戶關閉其他程式後重試'),
    ('Java 未安裝', 'ERR-JRE-001', '系統無 Java JRE 或版本低於 17', '引導至 Java 官方下載頁面'),
    ('麥克風無聲音', 'ERR-AUD-001', '麥克風權限未開啟或設備故障', '顯示麥克風設定引導'),
    ('MIDI 播放無聲', 'ERR-MIDI-001', '系統 MIDI 設備未設定', '引導使用者至 Windows 音效設定'),
    ('磁碟空間不足', 'ERR-FS-001', '剩餘空間 < 500MB', '提示清理磁碟空間'),
    ('MIDI 設備被佔用', 'ERR-MIDI-002', '其他程式正在使用 MIDI 設備', '提示關閉其他音樂軟體後重試'),
]
table_err = doc.add_table(rows=len(error_handling), cols=4)
table_err.style = 'Table Grid'
for i, row_data in enumerate(error_handling):
    for j, cell_data in enumerate(row_data):
        table_err.rows[i].cells[j].text = cell_data

doc.add_heading('9.3 安裝發布方案', level=2)
doc.add_paragraph('練琴寶提供兩種安裝方式：')

doc.add_heading('9.3.1 系統需求規格', level=3)
doc.add_paragraph('最低需求：')
doc.add_paragraph('• 作業系統：Windows 10 (64-bit)')
doc.add_paragraph('• 處理器：Intel Core i3 或同等 AMD 處理器')
doc.add_paragraph('• 記憶體：4 GB RAM')
doc.add_paragraph('• 磁碟空間：1 GB（含 Python、Java JRE、Audiveris）')
doc.add_paragraph('• 螢幕解析度：1280×720')
doc.add_paragraph('建議需求：')
doc.add_paragraph('• 作業系統：Windows 11 (64-bit)')
doc.add_paragraph('• 處理器：Intel Core i5 或同等 AMD 處理器')
doc.add_paragraph('• 記憶體：8 GB RAM')
doc.add_paragraph('• 磁碟空間：2 GB')
doc.add_paragraph('• 螢幕解析度：1920×1080')

doc.add_heading('9.3.2 相依套件清單', level=3)
doc.add_paragraph('Python 套件（版本為最低要求）：')
doc.add_paragraph('• PyQt6 >= 6.5.0（GUI 框架）')
doc.add_paragraph('• PyMuPDF >= 1.23.0（PDF 轉圖片）')
doc.add_paragraph('• OpenCV >= 4.8.0（影像前處理）')
doc.add_paragraph('• numpy >= 1.24.0（影像矩陣運算）')
doc.add_paragraph('• music21 >= 9.1.0（MusicXML 解析）')
doc.add_paragraph('• mido >= 1.3.0（MIDI 處理）')
doc.add_paragraph('• pygame >= 2.5.0（MIDI 播放）')
doc.add_paragraph('• Pillow >= 10.0.0（圖片處理）')
doc.add_paragraph('外部相依：')
doc.add_paragraph('• Java JRE 17+（Audiveris 運行環境）')
doc.add_paragraph('• Audiveris 6.3+（OMR 引擎，含 CLI 模式）')

doc.add_heading('9.3.3 MSI 安裝包（推薦）', level=3)
doc.add_paragraph('使用 WiX Toolset 或 NSIS 打包：')
doc.add_paragraph('• 自動安裝所有相依套件（Python、Java JRE）')
doc.add_paragraph('• 建立開始功能表捷徑與桌面捷徑')
doc.add_paragraph('• 支援 Windows 卸載程式完整移除')
doc.add_paragraph('• 支援 Silent Install（無聲安裝，適合大量部署）')
doc.add_paragraph('• Java JRE 安裝策略：安裝程式偵測系統是否已安裝 JRE 17+，若無則自動下載並安裝（內含 JRE 的精簡版，約 80MB）；若使用者拒絕自動安裝，則引導至 Java 官方網站下載')
doc.add_paragraph('• Audiveris 會跟隨安裝包一起部署，不依賴系統層級安裝')

doc.add_heading('9.3.4 PyInstaller 單一執行檔', level=3)
doc.add_paragraph('適用於快速分發、免安裝情境：')
doc.add_paragraph('• 打包為單一 .exe 檔案（約 200-300MB）')
doc.add_paragraph('• 包含 Python runtime、PyQt、Audiveris JRE')
doc.add_paragraph('• 缺點：檔案較大、啟動速度較慢、不支援獨立更新套件')

doc.add_heading('9.3.5 版本號規範', level=3)
doc.add_paragraph('採用 Semantic Versioning（major.minor.patch）：')
doc.add_paragraph('• Major：不相容的 API/資料格式變更（如樂譜檔案格式改版）')
doc.add_paragraph('• Minor：新增功能但向後相容（如新增調音器）')
doc.add_paragraph('• Patch：錯誤修正（如 OMR 辨識修正）')
doc.add_paragraph('• 版本號顯示於「關於」頁面，並寫入程式碼與安裝包檔案名稱')
doc.add_paragraph('• 格式範例：MusicMate_1.2.0_Setup.exe')

doc.add_heading('9.3.6 自動更新機制', level=3)
doc.add_paragraph('• 啟動時檢查 GitHub Releases 或自建伺服器最新版本')
doc.add_paragraph('• 若有新版本，彈出通知對話框：版本號 + 更新內容摘要 + 下載按鈕')
doc.add_paragraph('• 用戶可選擇「稍後提醒」或「立即下載」')
doc.add_paragraph('• 下載完成後自動執行安裝包，關閉舊版本程序')
doc.add_paragraph('• 檢查頻率：每次啟動檢查一次，非強制更新')

doc.add_heading('9.3.7 數位簽章', level=3)
doc.add_paragraph('• 發布前對 .exe / MSI 進行數位簽章（Authenticode）')
doc.add_paragraph('• 避免 Windows SmartScreen 警告')
doc.add_paragraph('• 簽章憑證需定期更新，建議使用 EV Code Signing 憑證')

doc.add_heading('9.3.8 CI/CD 流程', level=3)
doc.add_paragraph('• GitHub Actions 自動化打包流程：')
doc.add_paragraph('  1. Push to main branch → 觸發 CI 流程')
doc.add_paragraph('  2. 執行單元測試、Lint 檢查')
doc.add_paragraph('  3. 若測試通過，執行 PyInstaller 打包 + MSI 打包')
doc.add_paragraph('  4. 上傳建置成品至 GitHub Releases')
doc.add_paragraph('• 手動觸發：透過 GitHub Release 頁面上傳手動打包的檔案')

doc.add_heading('9.3.9 發布目錄結構', level=3)
doc.add_paragraph('PyInstaller 打包後的目錄結構：')
doc.add_paragraph('MusicMate/')
doc.add_paragraph('├── MusicMate.exe              # 主程式')
doc.add_paragraph('├── _internal/                  # PyQt6、numpy 等 Python 套件')
doc.add_paragraph('├── audiveris/                  # Audiveris CLI + JRE')
doc.add_paragraph('│   ├── bin/audiveris-cli.bat')
doc.add_paragraph('│   └── jre/')
doc.add_paragraph('├── config/                     # 使用者設定檔（首次執行自動建立）')
doc.add_paragraph('├── logs/                       # 日誌目錄')
doc.add_paragraph('└── README.txt                  # 簡易說明')

doc.add_heading('9.3.10 發布檢查清單', level=3)
doc.add_paragraph('□ 所有單元測試通過（覆蓋率 > 80%）')
doc.add_paragraph('□ 在乾淨的 Windows 10/11 環境測試安裝')
doc.add_paragraph('□ 說明文件（使用手冊、FAQ）已完成')
doc.add_paragraph('□ 數位簽章完成')
doc.add_paragraph('□ Release Notes 已撰寫並附於安裝包與 GitHub Releases')
doc.add_paragraph('□ 自動更新伺服器版本號已更新')

doc.add_heading('9.3.11 崩潰自動復原', level=3)
doc.add_paragraph('• 應用程式異常崩潰時，下次啟動檢測未正常關閉旗標')
doc.add_paragraph('• 若檢測到崩潰，顯示復原對話框：')
doc.add_paragraph('  - 詢問用戶是否恢復上次開啟的樂譜')
doc.add_paragraph('  - 提供「查看崩潰日誌」按鈕，便於回報問題')
doc.add_paragraph('• 正常關閉時清除旗標，避免誤觸發')
doc.add_paragraph('• 進行中的 OMR 任務若未完成，清除暫存檔避免殘留')
doc.add_paragraph('• 設定檔採用原子寫入（先寫入暫存檔再 rename），避免寫入中斷導致損壞')

doc.add_page_break()

# ==================== SECTION 10: MILESTONES ====================
doc.add_heading('10. 開發里程碑（v1.0）', level=1)

milestones = [
    ('階段', '內容', '優先序', '預估工期', '前置依賴'),
    ('Phase 0', '專案初始化、架構設計、環境建置', '★★★★★', '1 週', '—'),
    ('Phase 1', 'GUI 基本框架、Tab 導航架構（4 Tab）', '★★★★★', '2 週', 'Phase 0'),
    ('Phase 2', '節拍器功能實作', '★★★★☆', '1 週', 'Phase 1'),
    ('Phase 3', '調音器功能實作', '★★★★☆', '1 週', 'Phase 1'),
    ('Phase 4', 'OMR 系統整合（PyMuPDF + Audiveris + music21）', '★★★★★', '3 週', 'Phase 1'),
    ('Phase 5', '深度學習音符辨識模型訓練', '★★★☆☆', ' TBD', 'Phase 4 完成後評估'),
    ('Phase 6', '自動翻譜頁整合 OMR', '★★★★☆', '2 週', 'Phase 4'),
    ('Phase 7', 'MIDI 播放功能實作（mido + pygame）', '★★★★☆', '2 週', 'Phase 4'),
    ('Phase 8', '樂譜播放頁與自動翻譜連動', '★★★☆☆', '1 週', 'Phase 6 + Phase 7'),
    ('Phase 9', '節拍器與播放同步連動', '★★★☆☆', '1 週', 'Phase 2 + Phase 7'),
    ('Phase 10', '測試、除錯、優化', '★★★★☆', '2 週', 'Phase 8 + Phase 9'),
    ('Phase 11', '第一版發布（打包、簽章、CI/CD）', '★★★★★', '1 週', 'Phase 10'),
]
table6 = doc.add_table(rows=len(milestones), cols=5)
table6.style = 'Table Grid'
for i, row_data in enumerate(milestones):
    for j, cell_data in enumerate(row_data):
        table6.rows[i].cells[j].text = cell_data

doc.add_heading('10.1 平行開發建議', level=2)
doc.add_paragraph('• Phase 2（節拍器）與 Phase 3（調音器）可平行開發，互不依賴')
doc.add_paragraph('• Phase 4（OMR）與 Phase 2/3 可平行開發')
doc.add_paragraph('• Phase 6（翻譜整合）與 Phase 7（MIDI 播放）可平行開發，兩者皆依賴 Phase 4')
doc.add_paragraph('• Phase 8 與 Phase 9 需等待前述階段完成')

doc.add_heading('10.2 總工期估算', level=2)
doc.add_paragraph('• 關鍵路徑：Phase 0 → Phase 1 → Phase 4 → Phase 6/7 → Phase 8 → Phase 10 → Phase 11')
doc.add_paragraph('• 估算總工期：約 14-16 週（單一開發者）')
doc.add_paragraph('• 若有兩人開發，可平行處理 Phase 2/3 與 Phase 4，縮短至約 12 週')

doc.add_page_break()

# ==================== SECTION 11: CHANGE LOG ====================
doc.add_heading('11. 版本變更紀錄', level=1)

change_log = [
    ('版本', '日期', '變更內容'),
    ('v0.1', '2026-04-19', '初始版本：產品功能列表、技術架構、UI 規格'),
    ('v0.2', '2026-04-19', '新增 UI 規格詳細設計'),
    ('v0.3', '2026-04-20', '整合 Chat Log，統一格式'),
    ('v0.4', '2026-04-21', 'OMR 改用 Audiveris（Phase 4），新增 PDF OMR 處理流程（Section 7），Phase 5 調整為未來規劃'),
    ('v0.5', '2026-04-21', '新增 Section 8 非功能需求（效能、相容性、可靠度、可用性、安全性、擴展性）'),
    ('v0.6', '2026-04-21', 'Section 8 擴充：新增擴展性細節，調整 Section 9 錯誤處理與安裝發布'),
    ('v0.7', '2026-04-21', 'Section 9 微調：修正錯字、完善表格格式'),
    ('v0.8', '2026-04-27', 'Section 9 大幅擴充：日誌規範、自訂 Exception、Audiveris 崩潰處理、錯誤對話框、更多錯誤情境、系統需求、相依套件、版本號規範、自動更新、數位簽章、CI/CD、發布目錄結構'),
    ('v0.9', '2026-04-27', '全面審查：新增 Section 1 產品概述、Section 4 系統資料流程、修複 Section 3 不一致、新增 Tap Tempo/音高演算法/循環播放/移調/分段練習、新增設定對話框/快捷鍵表/關於對話框/拖放支援/檔案選單、新增 Section 8.7 資料持久化/8.8 檔案關聯、新增 9.3.11 崩潰自動復原、Section 10 新增工期與依賴關係、新增 Section 11 版本紀錄、新增附錄名詞解釋'),
    ('v1.0', '2026-04-29', '規格書最終審查版：明確 A4 預設值（440 Hz）與適用情境、新增 Soundfont 音源方案說明、新增多音軌樂譜音軌選擇 UI、新增自動翻譜同步機制（跨頁長音符/空白頁/提前量）、新增測試策略（pytest、覆蓋率目標 80%）、更新 MSI 安裝策略（Java JRE 自動偵測與隨包部署）、更新技術堆疊（sounddevice、預錄 WAV 樣本）、更新版本紀錄與附錄'),
    ('v1.1', '2026-04-29', '新增 Section 4.6 節拍器流程（完整操作邏輯：強弱拍/音效來源/TapTempo/計時器）、節拍器音色擴充至6種（木魚/小鼓/叮叮聲/狗吠等）、預設使用主機內建音效合成器、WAV 樣本改為備援'),
]
table_cl = doc.add_table(rows=len(change_log), cols=3)
table_cl.style = 'Table Grid'
for i, row_data in enumerate(change_log):
    for j, cell_data in enumerate(row_data):
        table_cl.rows[i].cells[j].text = cell_data

doc.add_page_break()

# ==================== APPENDIX: GLOSSARY ====================
doc.add_heading('附錄：名詞解釋', level=1)

glossary = [
    ('名詞', '說明'),
    ('OMR', 'Optical Music Recognition，光學樂譜辨識，將掃描的樂譜圖片轉換為結構化音樂資料'),
    ('MusicXML', '標準的音樂交換格式，用於在不同音樂軟體之間交換樂譜資料'),
    ('MIDI', 'Musical Instrument Digital Interface，數位樂器通訊標準，用於電子樂器之間交換演奏資訊'),
    ('General MIDI（GM）', '標準化的 MIDI 音色映射規範，確保不同設備播放同一 MIDI 檔案時音色一致'),
    ('Cents', '音分，音高單位，1 個半音 = 100 cents，用於精確描述音高偏差'),
    ('A440', '標準音高，A4 音符頻率為 440 Hz，作為樂器調音基準'),
    ('BPM', 'Beats Per Minute，每分鐘節拍數，用於表示音樂速度'),
    ('FFT', 'Fast Fourier Transform，快速傅立葉變換，將音訊訊號從時域轉換為頻域，用於音高偵測'),
    ('自相關（Autocorrelation）', '信號處理演算法，計算信號與其延遲版本的相似度，常用於音高偵測'),
    ('PyQt6', 'Python 的 Qt6 GUI 框架綁定，用於建立跨平台桌面應用程式'),
    ('OpenCV', 'Open Source Computer Vision Library，開源電腦視覺函式庫'),
    ('Audiveris', '開源 OMR 引擎，Java 實作，支援命令列批次處理'),
    ('PyMuPDF（fitz）', 'Python 的 PDF 處理函式庫，基於 MuPDF，支援 PDF 轉圖片'),
    ('music21', 'Python 的音樂學分析函式庫，支援 MusicXML 解析與生成'),
    ('mido', 'Python 的 MIDI 處理函式庫，用於解析與生成 MIDI 訊息'),
    ('PyInstaller', 'Python 程式打包工具，可將程式打包為單一可執行檔'),
    ('WiX Toolset', 'Windows Installer XML，用於建立 MSI 安裝包的開源工具'),
    ('Semantic Versioning', '語意化版本號規範，格式為 major.minor.patch'),
    ('CI/CD', 'Continuous Integration / Continuous Deployment，持續整合/持續部署'),
    ('Authenticode', 'Microsoft 的數位簽章技術，用於驗證軟體來源與完整性'),
]
table_gl = doc.add_table(rows=len(glossary), cols=2)
table_gl.style = 'Table Grid'
for i, row_data in enumerate(glossary):
    for j, cell_data in enumerate(row_data):
        table_gl.rows[i].cells[j].text = cell_data

doc.add_page_break()

p = doc.add_paragraph()
p.add_run('本文件為產品規格書 v1.0，待 William 審閱確認後作為軟體工程師實作依據。').italic = True
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('※ v1.1 更新（2026-04-29）：新增 Section 4.6 節拍器流程、節拍器音色擴充至6種（木魚/小鼓/叮叮聲/狗吠等）、預設主機內建音效合成器、附四個功能頁面完整 UI 示意圖。').italic = True

output_path = '/home/aping/.openclaw/workspace/projects/musicmate/docs/MusicMate_Product_Spec_v1.1.docx'
doc.save(output_path)
print(f'Full specification document saved: {output_path}')
