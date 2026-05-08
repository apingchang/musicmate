#!/usr/bin/env python3
"""Generate comprehensive Word document for MusicMate full specification v0.3."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 設定頁面邊距（單位：公分）
# 左、右、上、下 = 1.00 cm，裝訂邊 = 0.5 cm
section = doc.sections[0]
section.left_margin = Cm(1.00)
section.right_margin = Cm(1.00)
section.top_margin = Cm(1.00)
section.bottom_margin = Cm(1.00)
section.gutter = Cm(0.5)

# ==================== TITLE PAGE ====================
title = doc.add_heading('練琴寶 / MusicMate', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('完整產品規格書', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('版本：v0.4 Draft').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('日期：2026-04-21')
doc.add_paragraph('負責人：William（客人）/ 夥計（協調主持）')
doc.add_paragraph()
doc.add_paragraph('※ 本文件包含練琴寶所有已討論過的產品規格，待 William 審閱後修改。')
doc.add_page_break()

# ==================== SECTION 2: FEATURES ====================
doc.add_heading('2. 產品功能列表', level=1)

doc.add_heading('2.1 自動翻譜功能', level=2)
features1 = [
    ('功能描述', '接收 PDF 或圖片格式的樂譜，自動顯示並根據樂曲速度自動滾動'),
    ('樂譜輸入', 'PDF 檔案開啟、圖片匯入（ JPG/PNG）、相機拍攝'),
    ('樂譜辨識', '使用 OMR（光學樂譜辨識）技術自動分析音符、節拍、速度'),
    ('顯示功能', '支援滑鼠滾輪手動滾動、雙擊全螢幕模式'),
    ('控制功能', '播放/暫停、上一頁/下一頁、速度調整（0.5x ~ 3.0x）'),
    ('快捷鍵', 'Space=播放/暫停、←→=換頁、F=全螢幕'),
    ('特殊功能', '無樂譜時顯示引導區塊、支援拖放檔案'),
]
table1 = doc.add_table(rows=len(features1), cols=2)
table1.style = 'Table Grid'
for i, (k, v) in enumerate(features1):
    table1.rows[i].cells[0].text = k
    table1.rows[i].cells[1].text = v

doc.add_heading('2.2 節拍器功能', level=2)
features2 = [
    ('功能描述', '提供精準節拍，協助練琴時維持穩定速度'),
    ('BPM 範圍', '40 ~ 240 BPM'),
    ('BPM 控制', '-10/-1/+1/+10 微調按鈕、快捷 60/80/100/120'),
    ('拍號選擇', '2/4、3/4、4/4、5/4、6/8、7/8、12/8'),
    ('節拍音效', 'Click / Woodblock / Digital 三種音色'),
    ('音量控制', '獨立即量滑桿'),
    ('計時器', '5/10/15/30/45/60 分鐘倒數計時'),
    ('視覺效果', '節拍燈閃爍動畫'),
]
table2 = doc.add_table(rows=len(features2), cols=2)
table2.style = 'Table Grid'
for i, (k, v) in enumerate(features2):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v

doc.add_heading('2.3 調音器功能', level=2)
features3 = [
    ('功能描述', '收音麥克風偵測音高，協助樂器對音'),
    ('顯示方式', '半圓弧形指針儀表'),
    ('音高顯示', '音名（C/C#/D...）、Cents 偏差值、Hz 頻率'),
    ('精準範圍', '±5 cents 綠色區間 = 準確'),
    ('麥克風控制', '點擊啟動/停用，收音時顯示綠光指示'),
    ('參考音設定', 'A4 頻率（預設 440 Hz，可調 430~450 Hz）'),
    ('快捷參考音', '432 Hz / 440 Hz / 442 Hz'),
    ('樂器預設', '吉他/烏克麗麗/小提琴/大提琴/長笛等'),
]
table3 = doc.add_table(rows=len(features3), cols=2)
table3.style = 'Table Grid'
for i, (k, v) in enumerate(features3):
    table3.rows[i].cells[0].text = k
    table3.rows[i].cells[1].text = v

doc.add_heading('2.4 樂譜播放功能（Audio Playback）【新】', level=2)
features4 = [
    ('功能描述', '使用電腦 MIDI 音源播放樂譜，邊聽邊練琴'),
    ('輸入來源', 'OMR 辨識後的音符結構'),
    ('MIDI 合成', '使用 General MIDI 內建模音源播放'),
    ('音色選擇', 'GM 標準音色（鋼琴、吉他、小提琴、長笛等）'),
    ('速度控制', '0.5x ~ 2.0x 播放速度調整'),
    ('連動功能', '邊播放邊自動翻譜（與自動翻譜頁連動）'),
    ('播放控制', '播放/暫停/停止、進度條顯示'),
    ('當前音符', '高亮顯示目前播放的音符位置'),
]
table4 = doc.add_table(rows=len(features4), cols=2)
table4.style = 'Table Grid'
for i, (k, v) in enumerate(features4):
    table4.rows[i].cells[0].text = k
    table4.rows[i].cells[1].text = v

doc.add_page_break()

# ==================== SECTION 3: TECH ARCHITECTURE ====================
doc.add_heading('3. 技術架構', level=1)

doc.add_heading('3.1 系統架構', level=2)
doc.add_paragraph('練琴寶採用三層架構設計：')
doc.add_paragraph('• 表現層（Presentation Layer）：Python GUI（PyQt/PySide）')
doc.add_paragraph('• 業務邏輯層（Business Logic）：核心功能模組')
doc.add_paragraph('• 資料處理層（Data Layer）：OMR 引擎、檔案處理')

doc.add_heading('3.2 OMR 光學樂譜辨識流程', level=2)
doc.add_paragraph('1. 樂譜圖片輸入（PDF 轉圖片 or 直接圖片）')
doc.add_paragraph('2. OpenCV 前處理：灰階轉換 → 降噪 → 傾斜校正 → 二值化')
doc.add_paragraph('3. 音符區域切割')
doc.add_paragraph('4. 深度學習模型辨識音符與節拍資訊')
doc.add_paragraph('5. 樂理結構化輸出（音符、時值、節拍、速度）')

doc.add_heading('3.3 MIDI 播放流程', level=2)
doc.add_paragraph('1. OMR 輸出（JSON 音符結構）→ mido 處理')
doc.add_paragraph('2. 轉換為 MIDI 訊息（Note On/Off、Program Change）')
doc.add_paragraph('3. pygame.simpleaudio 播放 MIDI 音頻')
doc.add_paragraph('4. 選擇 GM 音色（Program Change訊息）')
doc.add_paragraph('5. 速度調整（Playback Rate）')

doc.add_heading('3.4 主要技術堆疊', level=2)
tech_stack = [
    ('項目', '技術選擇'),
    ('程式語言', 'Python 3.12+'),
    ('GUI 框架', 'PyQt6 或 PySide6（Windows Fluent Design）'),
    ('影像處理', 'OpenCV (cv2)'),
    ('OMR 辨識', '深度學習模型（自訓練或預訓練模型微調）'),
    ('音頻處理', 'numpy、scipy（音訊分析）'),
    ('節拍器音效', 'pygame 或 winsound'),
    ('MIDI 處理', 'mido（MIDI 訊息解析與生成）'),
    ('音訊播放', 'pygame / simpleaudio'),
    ('調音器收音', 'pyaudio'),
    ('PDF 處理', 'PyMuPDF（fitz）或 pdf2image'),
    ('Word 輸出', 'python-docx'),
]
table5 = doc.add_table(rows=len(tech_stack), cols=2)
table5.style = 'Table Grid'
for i, (k, v) in enumerate(tech_stack):
    table5.rows[i].cells[0].text = k
    table5.rows[i].cells[1].text = v

doc.add_page_break()

# ==================== SECTION 5: UI ====================
doc.add_heading('5. Windows 版 UI 規格', level=1)

doc.add_heading('5.1 整體架構', level=2)
doc.add_paragraph('• 視窗模型：單一主視窗 + 底部 Tab 導航（4個 Tab）')
doc.add_paragraph('• 最小尺寸：800 × 600 px')
doc.add_paragraph('• 預設尺寸：1200 × 800 px')
doc.add_paragraph('• 可縮放：是（響應式佈局）')
doc.add_paragraph('• 標題列：標準 Windows 標題列')
doc.add_paragraph('• 主題：跟隨 Windows 系統深色/淺色設定')

doc.add_heading('5.2 佈局結構', level=2)
doc.add_paragraph('單一視窗分為以下四個區塊（由上而下）：')
doc.add_paragraph()

# 插入 UI 佈局示意圖
diagram_path = '/home/aping/.openclaw/workspace/projects/musicmate/docs/ui_layout_diagram.png'
doc.add_picture(diagram_path, width=Inches(6.0))
doc.add_paragraph()
doc.add_paragraph('▲ 圖：練琴寶 MusicMate Windows 版 UI 佈局結構')

doc.add_heading('5.3 四大功能頁面', level=2)

# Page 1: Auto-scroll
doc.add_heading('5.3.1 自動翻譜頁', level=3)
doc.add_paragraph('用途：顯示樂譜 PDF/圖片，自動向上滾動，充當虛擬樂譜架。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 樂譜顯示區：PDF/圖片顯示、支援滑鼠滾輪、雙擊全螢幕')
doc.add_paragraph('• 控制工具列：開啟樂譜、播放/暫停、速度滑桿（0.5x~3.0x）')
doc.add_paragraph('• 頁面導航：上一頁/下一頁按鈕')
doc.add_paragraph('• 快捷鍵：Space=播放/暫停、←→=換頁、F=全螢幕')

# Page 2: Metronome
doc.add_heading('5.3.2 節拍器頁', level=3)
doc.add_paragraph('用途：精準節拍器，協助練琴時維持穩定速度。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 節拍顯示：BPM 大字體、拍號、閃爍動畫')
doc.add_paragraph('• BPM 控制：40~240、快捷 60/80/100/120')
doc.add_paragraph('• 拍號選擇：2/4、3/4、4/4、5/4、6/8、7/8、12/8')
doc.add_paragraph('• 音效：Click/Woodblock/Digital')
doc.add_paragraph('• 計時器：5~60 分鐘倒數')

# Page 3: Tuner
doc.add_heading('5.3.3 調音器頁', level=3)
doc.add_paragraph('用途：收音麥克風偵測音高，協助樂器對音。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 指針顯示：半圓弧形，±5 cents 綠色=準')
doc.add_paragraph('• 音高顯示：音名、Cents、Hz')
doc.add_paragraph('• 麥克風：啟動/停用控制')
doc.add_paragraph('• 參考音：432/440/442 Hz 快捷鍵')
doc.add_paragraph('• 樂器預設：吉他/小提琴/長笛等')

# Page 4: Audio Playback【新】
doc.add_heading('5.3.4 樂譜播放頁（Audio Playback）【新】', level=3)
doc.add_paragraph('用途：使用電腦 MIDI 音源播放樂譜，邊聽邊練琴。')
doc.add_paragraph('子區塊：')
doc.add_paragraph('• 播放控制：播放/暫停、停止按鈕')
doc.add_paragraph('• 速度控制：滑桿 0.5x ~ 2.0x')
doc.add_paragraph('• 音色選擇：下拉選單（鋼琴/吉他/小提琴等 GM 音色）')
doc.add_paragraph('• 音量控制：獨立即量滑桿')
doc.add_paragraph('• 進度條：顯示目前播放位置')
doc.add_paragraph('• 當前音符：高亮顯示目前播放的音符')

doc.add_heading('5.3.5 連動功能說明', level=3)
doc.add_paragraph('自動翻譜頁 + 樂譜播放頁可同時運行：')
doc.add_paragraph('• 啟動播放時，樂譜自動跟著滾動')
doc.add_paragraph('• 節拍器可與播放同步')
doc.add_paragraph('• 可隨時切換手動/自動模式')

doc.add_page_break()

# ==================== SECTION 6: OMR ====================
doc.add_heading('6. 光學樂譜辨識（OMR）方案', level=1)

doc.add_heading('6.1 為何使用 OpenCV？', level=2)
doc.add_paragraph('• Python 原生，跨平台方便')
doc.add_paragraph('• 影像前處理功能強大（降噪、傾斜校正、二值化）')
doc.add_paragraph('• 生態系完整，搭配其他函式庫容易')

doc.add_heading('6.2 OMR 處理流程（Prototype → Audiveris）', level=2)
doc.add_paragraph('第一版採用 Audiveris（開源 OMR）做 Prototype，流程如下：')
doc.add_paragraph('Step 1：樂譜輸入（PDF 轉圖片 or 直接圖片）')
doc.add_paragraph('Step 2：PDF 處理（PyMuPDF 轉圖片）')
doc.add_paragraph('Step 3：呼叫 Audiveris 命令列介面（CLI）進行樂譜辨識')
doc.add_paragraph('Step 4：輸出 MusicXML 格式音符結構')
doc.add_paragraph('Step 5：解析 MusicXML，轉換為內部 JSON 結構')

doc.add_heading('6.3 Audiveris 整合方案', level=2)
doc.add_paragraph('Audiveris 是目前最成熟的開源 OMR，採用以下整合方式：')
doc.add_paragraph('• 使用 N4 (N Audrey) 版本，支援命令列批次處理')
doc.add_paragraph('• Java 環境：需安裝 JRE 17+（Windows 11 環境）')
doc.add_paragraph('• Python 整合：subprocess 呼叫 Audiveris CLI，讀取 MusicXML 輸出')
doc.add_paragraph('• PDF 轉圖片：使用 PyMuPDF（fitz）處理 PDF 頁面轉圖片')
doc.add_paragraph('• 支援格式：PDF / JPG / PNG / TIFF 輸入')
doc.add_paragraph('• 輸出格式：MusicXML（標準音樂交換格式）')

doc.add_heading('6.4 技術堆疊（OMR 部分）', level=2)
omr_stack = [
    ('項目', '技術選擇'),
    ('PDF 處理', 'PyMuPDF（fitz）'),
    ('OMR 引擎', 'Audiveris（開源，Java）'),
    ('OMR 前處理', 'OpenCV 輔助（傾斜校正、降噪）'),
    ('音樂格式解析', 'music21（讀取 MusicXML）'),
    ('內部資料結構', 'JSON（音符、時值、節拍、小節）'),
]
table_omr = doc.add_table(rows=len(omr_stack), cols=2)
table_omr.style = 'Table Grid'
for i, (k, v) in enumerate(omr_stack):
    table_omr.rows[i].cells[0].text = k
    table_omr.rows[i].cells[1].text = v

doc.add_heading('6.5 Phase 5 規劃（未來自訓練模型）', level=2)
doc.add_paragraph('第一版使用 Audiveris 快速驗證可行性。\n未來若要提升準確度、支援手寫樂譜，可考虑：')
doc.add_paragraph('• 採用公開資料集訓練（MUSCIMA++、HASYTAC）')
doc.add_paragraph('• 模型架構：CNN + CTC（主流 OMR 做法）')
doc.add_paragraph('• 評估時機：Audiveris 準確度瓶頸出现後再決定')
doc.add_paragraph('• 優先序：★★★☆☆（目前擺後面）')

doc.add_page_break()

# ==================== SECTION 7: PDF OMR PROCESS ====================
doc.add_heading('7. PDF 樂譜 OMR 處理流程（第一版）', level=1)

doc.add_heading('7.1 整體流程', level=2)
doc.add_paragraph('PDF 格式樂譜的 OMR 處理分為以下 6 個步驟：')
process_steps = [
    ('步驟', '內容', '技術工具'),
    ('Step 1', '使用者開啟 PDF 樂譜', 'PyQt6 檔案對話框'),
    ('Step 2', 'PDF 轉圖片（每頁一張）', 'PyMuPDF（fitz）'),
    ('Step 3', '圖片前處理（傾斜/降噪/二值化）', 'OpenCV'),
    ('Step 4', 'Audiveris OMR 辨識', 'subprocess 呼叫 Audiveris CLI'),
    ('Step 5', '輸出 MusicXML 檔案', 'Audiveris'),
    ('Step 6', '解析 MusicXML 轉 JSON', 'music21'),
]
table_proc = doc.add_table(rows=len(process_steps), cols=3)
table_proc.style = 'Table Grid'
for i, row_data in enumerate(process_steps):
    for j, cell_data in enumerate(row_data):
        table_proc.rows[i].cells[j].text = cell_data

doc.add_heading('7.2 詳細步驟說明', level=2)
doc.add_paragraph('【Step 1】使用者開啟 PDF 樂譜')
doc.add_paragraph('• 使用者點擊「開啟樂譜」按鈕，選擇 PDF 檔案')
doc.add_paragraph('• 練琴寶檢查檔案是否存在，記錄檔案路徑')
doc.add_paragraph('• 顯示載入進度條（PDF 頁數可能達到數十頁）')

doc.add_paragraph('【Step 2】PDF 轉圖片（PyMuPDF）')
doc.add_paragraph('• 使用 PyMuPDF（fitz）將 PDF 每頁轉為圖片')
doc.add_paragraph('• 建議解析度：300 DPI（印刷品質）')
doc.add_paragraph('• 圖片格式：PNG（無失真，適合 OMR）')
doc.add_paragraph('• 支援多頁同時處理（批次轉換）')
doc.add_paragraph('• 臨時圖片存放在記憶體或暫存資料夾，處理完畢後自動刪除')

doc.add_paragraph('【Step 3】圖片前處理（OpenCV）')
doc.add_paragraph('• 傾斜校正（Deskew）：偵測水平線，自動旋轉')
doc.add_paragraph('• 灰階轉換（Grayscale）：降低運算量')
doc.add_paragraph('• 對比度增強（Contrast）：讓音符更明顯')
doc.add_paragraph('• 降噪（Denoise）：移除灰塵/掃描雜點')
doc.add_paragraph('• 二值化（Binarize）：轉為黑白，提高辨識度')

doc.add_paragraph('【Step 4】Audiveris OMR 辨識')
doc.add_paragraph('• 呼叫 Audiveris N4 命令列介面（CLI）')
doc.add_paragraph('• 批次處理多張圖片')
doc.add_paragraph('• Java 環境需 JRE 17+（Windows 11 環境）')
doc.add_paragraph('• 處理時間視 PDF 頁數而定（約數秒至數十秒）')

doc.add_paragraph('【Step 5】輸出 MusicXML 檔案')
doc.add_paragraph('• Audiveris 輸出標準 MusicXML 格式')
doc.add_paragraph('• 包含完整音符、時值、節拍、小節、調號資訊')
doc.add_paragraph('• 樂器軌道、動態、歌詞等額外資訊（可選）')

doc.add_paragraph('【Step 6】解析 MusicXML 轉 JSON（music21）')
doc.add_paragraph('• 使用 music21 函式庫讀取 MusicXML')
doc.add_paragraph('• 轉換為練琴寶內部 JSON 結構')
doc.add_paragraph('• 結構包含：頁數、tempo、拍號、每頁音符陣列')

doc.add_heading('7.3 輸出 JSON 結構範例', level=2)
doc.add_paragraph('{')
doc.add_paragraph('  "pages": [{"pageIndex": 0, "width": 800, "height": 1100}],')
doc.add_paragraph('  "tempo": 120,')
doc.add_paragraph('  "timeSignature": "4/4",')
doc.add_paragraph('  "keySignature": "C major",')
doc.add_paragraph('  "notes": [')
doc.add_paragraph('    {"page": 1, "x": 100, "y": 200, "pitch": "C4", "duration": "quarter", "measure": 1},')
doc.add_paragraph('    {"page": 1, "x": 150, "y": 200, "pitch": "D4", "duration": "quarter", "measure": 1}')
doc.add_paragraph('  ]')
doc.add_paragraph('}')

doc.add_page_break()

# ==================== SECTION 9: MILESTONES ====================
doc.add_heading('9. 開發里程碑（v0.4）', level=1)

milestones = [
    ('階段', '內容', '優先序'),
    ('Phase 0', '專案初始化、架構設計、環境建置', '★★★★★'),
    ('Phase 1', 'GUI 基本框架、Tab 導航架構（4 Tab）', '★★★★★'),
    ('Phase 2', '節拍器功能實作', '★★★★☆'),
    ('Phase 3', '調音器功能實作', '★★★★☆'),
    ('Phase 4', 'OMR 系統整合（PyMuPDF + Audiveris + music21）', '★★★★★'),
    ('Phase 5', '深度學習音符辨識模型訓練', '★★★☆☆'),
    ('Phase 6', '自動翻譜頁整合 OMR', '★★★★☆'),
    ('Phase 7', 'MIDI 播放功能實作（mido + pygame）', '★★★★☆'),
    ('Phase 8', '樂譜播放頁與自動翻譜連動', '★★★☆☆'),
    ('Phase 9', '節拍器與播放同步連動', '★★★☆☆'),
    ('Phase 10', '測試、除錯、優化', '★★★★☆'),
    ('Phase 11', '第一版發布', '★★★★★'),
]
table6 = doc.add_table(rows=len(milestones), cols=3)
table6.style = 'Table Grid'
for i, row_data in enumerate(milestones):
    for j, cell_data in enumerate(row_data):
        table6.rows[i].cells[j].text = cell_data

doc.add_page_break()

p = doc.add_paragraph()
p.add_run('本文件為產品規格書 v0.4，待 William 審閱後修改。').italic = True
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('※ v0.4 更新：OMR 改用 Audiveris（Phase 4），新增 PDF OMR 處理流程（Section 7），Phase 5 調整為未來規劃。').italic = True

output_path = '/home/aping/.openclaw/workspace/projects/musicmate/docs/MusicMate_Product_Spec_v0.4.docx'
doc.save(output_path)
print(f'Full specification document saved: {output_path}')
