#!/usr/bin/env python3
"""Generate MusicMate Development Log Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 設定頁面邊距
section = doc.sections[0]
section.left_margin = Cm(1.00)
section.right_margin = Cm(1.00)
section.top_margin = Cm(1.00)
section.bottom_margin = Cm(1.00)
section.gutter = Cm(0.5)

# ==================== TITLE ====================
title = doc.add_heading('MusicMate 練琴寶', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('開發日誌', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('版本：v1.0').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('日期：2026-04-29')
doc.add_paragraph('負責人：William Chang / 夥計')
doc.add_paragraph()
doc.add_paragraph('※ 本文件記錄 MusicMate 專案開發過程中的所有重要討論與決定。')
doc.add_page_break()

# ==================== 2026-04-29 ====================
doc.add_heading('2026-04-29 開發日誌', level=1)

# Morning session
doc.add_heading('上午 — 規格書 v0.9 → v1.1 審查與更新', level=2)

doc.add_heading('v0.9 規格書 Review', level=3)
doc.add_paragraph('William 要求對 v0.9 規格書進行全面審查，發現以下問題：')

issues = [
    ('內容缺失', 'Section 2 產品功能列表、Section 9.2 常見錯誤與解決方案、Section 9.1.2 自訂 Exception 類別、Section 11 版本變更紀錄皆為空白'),
    ('說明不足', '調音器 A4 預設值未說明不同音樂類型慣例、Soundfont 方案未確定、無 OMR 結果驗證機制、節拍器聲音生成方式未說明'),
    ('同步機制不明', '自動翻譜與 MIDI 播放的同步邏輯說明不足（跨頁長音符、空白頁處理）'),
    ('測試策略', '缺少單元測試策略、Phase 定義不夠具體'),
    ('安裝策略', 'MSI 安裝時 Java/Audiveris 的處理不明確'),
]
for title_text, content in issues:
    p = doc.add_paragraph()
    p.add_run(f'• {title_text}：').bold = True
    p.add_run(content)

doc.add_heading('v1.0 產生', level=3)
doc.add_paragraph('針對 Review 結果，執行 gen_spec_docx.py 產生 v1.0，並補強：')
updates_v1 = [
    '新增 Section 8.4.1 測試策略（pytest、覆蓋率 >80%）',
    '新增 Soundfont 音源方案說明',
    '新增多音軌樂譜音軌選擇 UI 說明',
    '新增自動翻譜同步機制（跨頁長音符、空白頁、提前量 0~2 小節）',
    '更新技術堆疊（sounddevice、預錄 WAV 樣本）',
    '更新 MSI 安裝 Java 自動偵測與隨包部署策略',
]
for u in updates_v1:
    doc.add_paragraph(f'• {u}')

doc.add_heading('節拍器音色與流程更新', level=3)
doc.add_paragraph('William 反映 Section 4 中沒有節拍器的資料操作流程，且節拍器聲音除了 WAV 檔外，預設要使用主機內建音效合成器。')
doc.add_paragraph('更新內容：')
updates_metronome = [
    '新增 Section 4.6 節拍器流程（完整操作邏輯）',
    '節拍器音色從 3 種擴充至 6 種：Click（點擊）/ Woodblock（木魚）/ Digital（數位）/ Snare（小鼓）/ Bell（叮叮聲）/ Dog Bark（狗吠）',
    '音效預設使用主機內建 Windows MIDI Synthesizer / System Sounds API',
    'WAV 樣本改為備援選項',
    '新增 Tap Tempo 邏輯說明',
    '新增計時器倒數結束自動停止說明',
]
for u in updates_metronome:
    doc.add_paragraph(f'• {u}')

doc.add_heading('v1.1 產生（含 UI 示意圖）', level=3)
doc.add_paragraph('William 反映首頁版號沒改到，且 4 個功能頁面只有主 UI 有圖，其餘只有文字說明。')
doc.add_paragraph('更新內容：')
updates_v1_1 = [
    '封面版號從 v1.0 改為 v1.1',
    '使用 AI 圖像生成工具產生 4 個功能頁面的 UI mockup：',
    '  - gui_auto_scroll.png — 自動翻譜頁',
    '  - gui_metronome.png — 節拍器頁',
    '  - gui_tuner.png — 調音器頁',
    '  - gui_playback.png — 樂譜播放頁',
    '四張圖片插入 Section 5.3 對應章節',
    'docx 檔案從 87KB 增至 365KB（含 5 張圖）',
]
for u in updates_v1_1:
    doc.add_paragraph(f'• {u}')

doc.add_heading('William 提問：docx 內嵌圖片是否會打到 AI 模型限制？', level=3)
doc.add_paragraph('回覆：不會。docx 內嵌圖片不消耗對話 token；python-docx 讀取時只讀文字，圖片是二進制 blob 不進對話。William 決定先不壓縮。')

# Afternoon session
doc.add_heading('下午 — 開發前置準備與 GitHub 初始化', level=2)

doc.add_heading('開發環境討論', level=3)
doc.add_paragraph('William 說明環境架構：')
env_points = [
    'Windows 11 是 Host，PyCharm 直接在 Windows 端開發',
    'VirtualBox Ubuntu 透過 shared folder 掛載到 /home/aping/kivy_shared/',
    'code 存在 shared folder，Windows 和 Ubuntu 都能 access',
    'PyCharm 已安裝，可直接 access /home/aping/kivy_shared/',
]
for p_text in env_points:
    doc.add_paragraph(f'• {p_text}')

doc.add_heading('GitHub 設定', level=3)
doc.add_paragraph('William 申請 GitHub 帳號（https://github.com/apingchang），建立 musicmate repository。')
doc.add_paragraph('URL: https://github.com/apingchang/musicmate')

doc.add_heading('初始化 Git Repository', level=3)
doc.add_paragraph('在 /home/aping/kivy_shared/musicmate/ 建立完整專案結構並初始化 Git。')

doc.add_heading('Python 環境設定', level=3)
doc.add_paragraph('在 Ubuntu 端建立虛擬環境：')
doc.add_paragraph('python3 -m venv venv')
doc.add_paragraph('pip install -r requirements.txt')
doc.add_paragraph('Dependencies: PyQt6, PyMuPDF, opencv-python, numpy, music21, mido, pygame, Pillow, sounddevice, pytest, pytest-qt')

doc.add_heading('專案資料夾結構', level=3)
doc.add_paragraph('已建立的結構：')
structure_items = [
    'src/main.py — 主程式進入點',
    'src/ui/main_window.py — 主視窗 + 4 個 Tab',
    'src/ui/auto_scroll_page.py — 自動翻譜頁（UI 骨幹）',
    'src/ui/metronome_page.py — 節拍器頁（UI 骨幹）',
    'src/ui/tuner_page.py — 調音器頁（UI 骨幹）',
    'src/ui/playback_page.py — 樂譜播放頁（UI 骨幹）',
    'src/core/metronome.py — 節拍器引擎骨架',
    'src/core/tuner.py — 調音器引擎骨架',
    'src/core/omr.py — OMR 處理引擎骨架',
    'src/core/midi_player.py — MIDI 播放引擎骨架',
    'src/utils/logger.py — 日誌設定',
    'src/utils/settings.py — 設定檔管理',
    'requirements.txt — Python 依賴列表',
    '.gitignore / README.md / LICENSE',
]
for item in structure_items:
    doc.add_paragraph(f'• {item}')

doc.add_heading('重要技術決策（待實作驗證）', level=3)
decisions = [
    ('節拍器音效', '主機內建音效合成器（預設），WAV 樣本備援'),
    ('Soundfont', '使用 Windows 內建 GM Soundfont，不捆绑額外檔案'),
    ('Java JRE', 'MSI 安裝時自動偵測並隨包部署 Audiveris JRE'),
    ('GitHub push', '使用 Personal Access Token 認證（已完成）'),
    ('docx 圖片', '365KB 不影響模型，維持現狀'),
]
for key, val in decisions:
    p = doc.add_paragraph()
    p.add_run(f'• {key}：').bold = True
    p.add_run(val)

doc.add_heading('邊開發邊改規格', level=3)
doc.add_paragraph('William 確認：規格書可以邊 Coding 邊修改。原則：')
doc.add_paragraph('• 工程師開工前以 v1.1 為基準')
doc.add_paragraph('• 每完成一個 Phase 更新一次規格書並遞增版本號')
doc.add_paragraph('• 重大決策（换 OMR 引擎、砍功能、加功能）立即更新並記錄進當天日誌')
doc.add_paragraph('• 更新規格書需透過 gen_spec_docx.py 腳本，禁止直接編輯 .docx')

doc.add_page_break()

# ==================== VERSION HISTORY ====================
doc.add_heading('版本變更紀錄', level=1)

changelog = [
    ('版本', '日期', '變更內容'),
    ('v1.0', '2026-04-29', '規格書 v0.9 Review — 新增測試策略、Soundfont 方案、多音軌 UI、同步機制、MSI Java 部署、技術堆疊更新'),
    ('v1.1', '2026-04-29', '新增 Section 4.6 節拍器流程、6 種節拍器音色、主機內建音效為預設、附 4 個功能頁面 UI 示意圖、封面版號修正'),
]
table = doc.add_table(rows=len(changelog), cols=3)
table.style = 'Table Grid'
for i, row_data in enumerate(changelog):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('本文件為 MusicMate 練琴寶開發日誌 v1.0，記錄所有開發討論與決定。').italic = True

output_path = '/home/aping/kivy_shared/musicmate/docs/MusicMate_Development_Log_v1.0.docx'
doc.save(output_path)
print(f'Development log saved: {output_path}')